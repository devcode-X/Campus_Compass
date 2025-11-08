# # src/utils.py
# import os, re, sys, types
# from pathlib import Path
# import pdfplumber
# from docx import Document as DocxDocument
# from pdf2image import convert_from_path
# from transformers import pipeline
# from PIL import Image

# # === LangChain Compatibility Patches (for other imports) ===
# if 'langchain.docstore' not in sys.modules:
#     sys.modules['langchain.docstore'] = types.ModuleType('langchain.docstore')
# if 'langchain.docstore.document' not in sys.modules:
#     import langchain_core.documents as lcd
#     sys.modules['langchain.docstore.document'] = types.ModuleType('langchain.docstore.document')
#     sys.modules['langchain.docstore.document'].Document = lcd.Document
# if 'langchain.text_splitter' not in sys.modules:
#     import langchain_text_splitters as lts
#     sys.modules['langchain.text_splitter'] = types.ModuleType('langchain.text_splitter')
#     sys.modules['langchain.text_splitter'].RecursiveCharacterTextSplitter = lts.RecursiveCharacterTextSplitter
# # ===========================================================

# # === Initialize Microsoft TrOCR pipeline ===
# print("🔹 Loading Microsoft TrOCR model... This may take 15–20s on first run.")
# ocr_pipe = pipeline("image-to-text", model="microsoft/trocr-base-stage1")
# print("✅ TrOCR OCR pipeline initialized successfully.")

# def ocr_image(img_path: str) -> str:
#     """Extract text from an image using Microsoft TrOCR pipeline."""
#     try:
#         result = ocr_pipe(Image.open(img_path).convert("RGB"))
#         if isinstance(result, list) and len(result) > 0:
#             return result[0].get("generated_text", "").strip()
#         return ""
#     except Exception as e:
#         print(f"⚠️ OCR failed on {img_path}: {e}")
#         return ""

# def read_pdf(path: str) -> str:
#     """
#     Reads both text-based and scanned PDFs.
#     Uses TrOCR for pages without extractable text.
#     """
#     text = []
#     path = Path(path)
#     if not path.exists():
#         raise FileNotFoundError(f"❌ File not found: {path}")

#     with pdfplumber.open(path) as pdf:
#         for i, page in enumerate(pdf.pages, start=1):
#             extracted = page.extract_text()
#             if extracted and extracted.strip():
#                 text.append(extracted)
#             else:
#                 print(f"🔍 OCR fallback (TrOCR) on page {i} of {path.name}...")
#                 images = convert_from_path(str(path), first_page=i, last_page=i)
#                 for img in images:
#                     temp_path = Path("temp_page.png")
#                     img.save(temp_path)
#                     try:
#                         ocr_text = ocr_image(str(temp_path))
#                         if ocr_text:
#                             text.append(ocr_text)
#                     finally:
#                         temp_path.unlink(missing_ok=True)
#     return "\n".join(text)

# def read_docx(path: str) -> str:
#     """Reads text from Word documents."""
#     doc = DocxDocument(path)
#     return "\n".join(p.text for p in doc.paragraphs)

# def read_text(path: str) -> str:
#     """Reads plain text files."""
#     with open(path, "r", encoding="utf-8") as f:
#         return f.read()

# def clean_text(s: str) -> str:
#     """Cleans up redundant newlines and spaces."""
#     s = s.replace("\r", "\n")
#     s = re.sub(r"\n{3,}", "\n\n", s)
#     s = s.strip()
#     return s

# def list_data_files(data_dir: str):
#     """Lists all PDF, DOCX, and TXT files for ingestion."""
#     p = Path(data_dir)
#     exts = [".pdf", ".docx", ".txt"]
#     return [str(f) for f in p.glob("*") if f.suffix.lower() in exts]


# src/utils.py
# src/utils.py
import os, re, io
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument
from transformers import pipeline
import fitz  # PyMuPDF
from PIL import Image
import easyocr

# 🧩 Initialize OCR models
print("🧠 Initializing OCR engines (TrOCR + EasyOCR)...")
trocr_pipe = pipeline("image-to-text", model="microsoft/trocr-small-printed")
easyocr_reader = easyocr.Reader(['en'])
print("✅ TrOCR + EasyOCR ready.\n")

def read_pdf(path: str) -> str:
    """Extracts text from text-based or scanned PDFs using pdfplumber + TrOCR + EasyOCR fallback."""
    path = Path(path)
    text = []
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                extracted = page.extract_text()
                if extracted and extracted.strip():
                    text.append(extracted)
                else:
                    print(f"🔍 OCR fallback on page {i} of {path.name}...")
                    doc = fitz.open(path)
                    page = doc.load_page(i - 1)
                    pix = page.get_pixmap(dpi=200)
                    img_bytes = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")

                    # Try TrOCR first
                    trocr_text = trocr_pipe(img)[0]["generated_text"].strip()

                    # If TrOCR fails (empty or symbols), fallback to EasyOCR
                    if not trocr_text or trocr_text.strip("*") == "":
                        easy_text = "\n".join(easyocr_reader.readtext(img_bytes, detail=0))
                        print(f"⚙️ EasyOCR fallback used on page {i}.")
                        text.append(easy_text)
                    else:
                        text.append(trocr_text)

    except Exception as e:
        print(f"⚠️ PDF extraction failed ({e}). Running full OCR via EasyOCR...")
        doc = fitz.open(path)
        for i, page in enumerate(doc, start=1):
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            easy_text = "\n".join(easyocr_reader.readtext(img_bytes, detail=0))
            text.append(easy_text)

    full_text = "\n".join(text)
    print(f"\n✅ FINAL OCR OUTPUT ({path.name}):\n{'='*60}\n{full_text[:1000]}\n{'='*60}\n")
    return full_text


def read_docx(path: str) -> str:
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs)


def read_text(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def clean_text(s: str) -> str:
    s = s.replace("\r", "\n")
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def list_data_files(data_dir: str):
    """
    Returns a list of valid document file paths (.pdf, .docx, .txt)
    inside the given directory.
    """
    p = Path(data_dir)
    exts = [".pdf", ".docx", ".txt"]
    files = [str(f) for f in p.glob("*") if f.suffix.lower() in exts]
    
    if not files:
        print(f"⚠️ No data files found in {data_dir}")
    else:
        print(f"📂 Found {len(files)} data files in {data_dir}")
    
    return files

